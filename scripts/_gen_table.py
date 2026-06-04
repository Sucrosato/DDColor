"""Generate Word document with DINOv3 SDT architecture table."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
section = doc.sections[0]
section.left_margin = Cm(2)
section.right_margin = Cm(2)

doc.add_paragraph('DINOv3 SDT 架构细节表', style='Title')

rows = [
    ['模块', '子层', '参数', '输入形状', '输出形状'],
    ['输入', '—', '—', '—', '(B, 3, 256, 256)'],
    ['CenterPadding', '—', 'multiple=16', '(B, 3, H, W)', '(B, 3, H_pad, W_pad)'],
    ['Patch Embedding', 'Conv2d 3→384, k=16 s=16', '—', '(B, 3, 256, 256)', '(B, 261, 384)'],
    ['DINOv3 ViT-S', '×12 Transformer Blocks', 'd=384, h=6, FFN=1536', '(B, 261, 384)', '(B, 261, 384)×13'],
    ['Token 拆分', '提取 [2,5,8,11], 丢弃 Register', '—', '(B, 261, 384)', '4×{(B,384,16,16),(B,384)}'],
    ['WeightedFusion', 'CLS Readout: Lin 768→384', '×4', '(B, 256, 768)', '(B, 256, 384)'],
    ['', 'Projection: Lin 384→256', '×4', '(B, 256, 384)', '(B, 256, 256)'],
    ['', 'Softmax 加权融合', '4 权重', '4×(B,256,256)', '(B, 256, 256)'],
    ['Reshape', 'permute + reshape', '—', '(B, 256, 256)', '(B, 256, 16, 16)'],
    ['SpatialDetail\nEnhancer', 'DWConv 3×3, groups=256, 残差', '—', '(B, 256, 16, 16)', '(B, 256, 16, 16)'],
    ['ColorQuery\nBottleneck\n(可选)', '100 个 Color Queries', 'Q=100, d=256', '—', '(100, B, 256)'],
    ['', 'Cross-Attn (Q→空间) ×3', 'h=8', '(100,B,256)+(256,B,256)', '(100, B, 256)'],
    ['', 'Self-Attn (Q→Q) ×3', 'h=8', '(100, B, 256)', '(100, B, 256)'],
    ['', 'FFN 256→1024→256 ×3', '—', '(100, B, 256)', '(100, B, 256)'],
    ['', 'Mean Pool → 全局颜色', '—', '(B, 100, 256)', '(B, 256, 1, 1)'],
    ['', 'Concat+Conv1×1(512→256)+残差', '—', '(B, 512, 16, 16)', '(B, 256, 16, 16)'],
    ['DySample 阶段1', 'DySample ×2 → Conv+BN+ReLU', 's=2, lp, g=4', '(B,256,16,16)', '(B,256,32,32)'],
    ['', 'DySample ×2 → Conv+BN+ReLU', 's=2, lp, g=4', '(B,256,32,32)', '(B,256,64,64)'],
    ['', 'Conv3×3+BN+ReLU (refinement)', '—', '(B,256,64,64)', '(B,256,64,64)'],
    ['DySample 阶段2', 'DySample ×2 → Conv+BN+ReLU', 's=2, lp, g=4', '(B,256,64,64)', '(B,256,128,128)'],
    ['', 'DySample ×2 → Conv+BN+ReLU', 's=2, lp, g=4', '(B,256,128,128)', '(B,256,256,256)'],
    ['', 'Conv3×3+BN+ReLU (refinement)', '—', '(B,256,256,256)', '(B,256,256,256)'],
    ['Output Conv', 'Conv3×3 256→128, ReLU', '—', '(B,256,256,256)', '(B,128,256,256)'],
    ['', 'Conv3×3 128→32, ReLU', '—', '(B,128,256,256)', '(B,32,256,256)'],
    ['', 'Conv1×1 32→2', '—', '(B,32,256,256)', '(B,2,256,256)'],
    ['输出', 'ab 色度通道', '—', '—', '—'],
]

table = doc.add_table(rows=len(rows), cols=5, style='Table Grid')

# Track merges to avoid double-merging
for i, row_data in enumerate(rows):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if i == 0:
            run.bold = True

# Header center
for j in range(5):
    table.rows[0].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Merge first column by module
i = 1
while i < len(rows):
    if rows[i][0] != '':
        span = 1
        for k in range(i + 1, len(rows)):
            if rows[k][0] == '':
                span += 1
            else:
                break
        if span > 1:
            table.cell(i, 0).merge(table.cell(i + span - 1, 0))
            p = table.cell(i, 0).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(9)
            p.runs[0].bold = True
        i += span
    else:
        i += 1

note = doc.add_paragraph()
note.paragraph_format.space_before = Pt(12)
run = note.add_run('注：ColorQueryBottleneck 为可选模块（use_color_queries=True），若不启用则为 Identity 直接透传。B 为 batch size，形状以 256×256 输入为例。')
run.font.size = Pt(8)
run.italic = True

doc.save('notes/dinov3_sdt_table.docx')
print('Saved: notes/dinov3_sdt_table.docx')
